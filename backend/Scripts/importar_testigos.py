# -*- coding: utf-8 -*-
"""
importar_testigos.py
====================
Importa masivamente los archivos .doc/.docx de remitos de envío de
estándares/testigos al sistema LIMSS.

Extrae por cada documento:
  - Nombre del testigo (en negrita, después de "a saber:")
  - Producto N° (código del fabricante)
  - Lote
  - Vencimiento
  - IR N° (número de IR del ERP)
  - Cantidad remitida (stock inicial)
  - Conservación

El código interno se genera automático: TEST-NNN (correlativo en la BD).

Uso:
    python importar_testigos.py --carpeta "\\\\SERVIDOR\\carpeta"
    python importar_testigos.py --carpeta "\\\\SERVIDOR\\carpeta" --dry-run

Requiere: pip install python-docx pyodbc
"""

import re
import sys
import argparse
import logging
from datetime import datetime, date
from pathlib import Path

CONFIG = {
    "db_server":        "Lamarserver",
    "db_name":          "LIMSS",
    "db_user":          "limss_app",
    "db_password":      "Limss2024#",   # ← cambiar
    "db_driver":        "SQL Server",
    "id_usuario_carga": 1,
    "stock_minimo":     0.1,
    "unidad_medida":    "g",
}

MESES_ES = {
    "enero": 1, "ene": 1,
    "febrero": 2, "feb": 2,
    "marzo": 3, "mar": 3,
    "abril": 4, "abr": 4,
    "mayo": 5,
    "junio": 6, "jun": 6,
    "julio": 7, "jul": 7,
    "agosto": 8, "ago": 8,
    "septiembre": 9, "sep": 9, "sept": 9,
    "octubre": 10, "oct": 10,
    "noviembre": 11, "nov": 11,
    "diciembre": 12, "dic": 12,
}


def setup_logging(carpeta_log: Path):
    log_file = carpeta_log / f"importacion_testigos_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s  %(levelname)-8s  %(message)s",
        handlers=[
            logging.FileHandler(log_file, encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )
    return log_file


def limpiar(texto: str) -> str:
    if not texto:
        return ""
    texto = re.sub(r'\*+', '', texto)
    return " ".join(texto.strip().split())


def parsear_vencimiento(texto: str):
    """Convierte varios formatos de fecha a date. Retorna None si no puede."""
    texto = limpiar(texto).strip(".")
    # Normalizar separadores: "Abril/2028" → "Abril 2028"
    texto = texto.replace('/', ' ').replace('-', ' ')
    m = re.search(r'(\d{1,2})?\s*(?:de\s+)?(\w+)\s+(?:de\s+)?(\d{2,4})', texto, re.IGNORECASE)
    if m:
        dia = int(m.group(1)) if m.group(1) else 28
        mes_str = m.group(2).lower()
        anio = int(m.group(3))
        if anio < 100:
            anio += 2000  # "26" → 2026
        mes = MESES_ES.get(mes_str)
        if mes:
            try:
                return date(anio, mes, min(dia, 28))
            except ValueError:
                return date(anio, mes, 28)
    # Formato: "31 Diciembre 2026" → ya cubierto arriba
    # Formato: "30/04/2028" o "30/04/28"
    m2 = re.search(r'(\d{1,2})[/\-](\d{1,2})[/\-](\d{2,4})', texto)
    if m2:
        d, mo, y = int(m2.group(1)), int(m2.group(2)), int(m2.group(3))
        if y < 100:
            y += 2000
        try:
            return date(y, mo, d)
        except ValueError:
            pass
    return None


def parsear_cantidad(texto: str):
    """Extrae el número de la cantidad remitida. Retorna None si no puede."""
    if not texto or '---' in texto:
        return None
    m = re.search(r'(\d+(?:[.,]\d+)?)', texto)
    if m:
        return float(m.group(1).replace(',', '.'))
    return None


def parsear_documento(docx_path: Path) -> dict | None:
    try:
        from docx import Document
    except ImportError:
        logging.error("  python-docx no instalado.")
        return None

    try:
        doc = Document(docx_path)
    except Exception as e:
        logging.error(f"  No se pudo abrir {docx_path.name}: {e}")
        return None

    # Extraer todo el texto del cuerpo — solo la PRIMERA copia del documento
    # (los documentos tienen 2-3 copias: "Copia para X", "Copia para recibo", etc.)
    primera_copia_encontrada = False
    segunda_copia_encontrada = False
    parrafos_primera = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if re.search(r'Copia para', txt, re.IGNORECASE):
            if not primera_copia_encontrada:
                primera_copia_encontrada = True
                # Salteamos este párrafo "Copia para X" y seguimos cargando
                continue
            else:
                segunda_copia_encontrada = True
                break
        parrafos_primera.append(p)

    parrafos_a_usar = parrafos_primera if segunda_copia_encontrada else list(doc.paragraphs)

    lineas = [p.text.strip() for p in parrafos_a_usar if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lineas.append(cell.text.strip())
    texto_completo = "\n".join(lineas)

    # ── Nombre del testigo ──
    nombre = ""

    # Estrategia 0 (PRIORITARIA): "traceable NOMBRE_PRODUCTO" en negrita
    # Ej: "...adjunto el estándar de referencia secundario traceable PAMOATO DE PIRANTEL"
    for p in parrafos_a_usar:
        texto_bold = "".join(
            run.text for run in p.runs if run.bold and run.text.strip()
        ).strip()
        m_trace = re.search(r'traceable\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s\-DE]+?)(?:\s+Lote|\s+Producto|$)', texto_bold, re.IGNORECASE)
        if m_trace:
            nombre = limpiar(m_trace.group(1))
            break

    # Estrategia 1: buscar después de "a saber:" (solo si aún no tenemos nombre)
    if not nombre:
      en_nombre = False
      for p in parrafos_a_usar:
        # Caso A: "a saber:" al final de párrafo → nombre en siguiente párrafo
        if re.search(r'a saber\s*:\s*$', p.text.strip(), re.IGNORECASE):
            en_nombre = True
            continue
        # Caso B: "a saber:" en medio del párrafo → nombre puede estar en el bold del mismo párrafo
        if re.search(r'a saber\s*:', p.text, re.IGNORECASE) and not en_nombre:
            texto_bold = "".join(
                run.text for run in p.runs if run.bold and run.text.strip()
            ).strip()
            texto_bold = limpiar(texto_bold)
            # Buscar texto en mayúsculas dentro del bold
            m_inline = re.search(r'([A-ZÁÉÍÓÚÑ]{3,}(?:\s+[A-ZÁÉÍÓÚÑ]{2,})+)', texto_bold)
            if m_inline:
                candidato = limpiar(m_inline.group(1))
                if es_nombre_valido(candidato) and not re.search(r'Producto|Lote|N[°º]|Vencimiento|Cantidad|Conservaci|frasco|Copia', candidato, re.IGNORECASE):
                    nombre = candidato
                    break
            en_nombre = True
            continue
        if en_nombre:
            texto_bold = "".join(
                run.text for run in p.runs if run.bold and run.text.strip()
            ).strip()
            texto_bold = limpiar(texto_bold)
            if texto_bold and es_nombre_valido(texto_bold) and not re.search(r'Producto|Lote|N[°º]|Vencimiento|Cantidad|Conservaci|frasco|Copia|Saludo|atentamente', texto_bold, re.IGNORECASE):
                nombre = texto_bold
                break
            candidato = limpiar(p.text.strip())
            if candidato and es_nombre_valido(candidato) and not re.search(r'Producto|Lote|N[°º]|Vencimiento|Cantidad|Conservaci|frasco|Copia|Saludo|atentamente', candidato, re.IGNORECASE):
                nombre = candidato
                break

    # Estrategia 2: buscar línea en MAYÚSCULAS en negrita en el cuerpo
    if not nombre:
        for p in parrafos_a_usar:
            texto_bold = "".join(
                run.text for run in p.runs if run.bold and run.text.strip()
            ).strip()
            texto_bold = limpiar(texto_bold)
            if (texto_bold and texto_bold == texto_bold.upper()
                    and es_nombre_valido(texto_bold)
                    and not re.search(r'Producto|Lote|N[°º]|Vencimiento|Laboratorio|Especialidades|Copia|Cantidad|Conservaci', texto_bold, re.IGNORECASE)):
                nombre = texto_bold
                break

    # Estrategia 3: fallback con texto en mayúsculas
    if not nombre:
        idx = texto_completo.lower().find("a saber:")
        if idx >= 0:
            resto = texto_completo[idx+8:].strip().split('\n')
            for linea in resto:
                linea_clean = limpiar(linea)
                if linea_clean and linea_clean.isupper() and len(linea_clean) > 3:
                    nombre = linea_clean
                    break

    # ── Producto N° ──
    m = re.search(r'Producto\s+N[°º]?\s*[:.]?\s*(.+?)(?:\n|$)', texto_completo, re.IGNORECASE)
    nro_producto = limpiar(m.group(1)) if m else ""

    # ── Lote (también acepta Control Nº) ──
    m = re.search(r'(?:Lote|Control)\s+N[°º]?\s*[:.]?\s*(.+?)(?:\n|$)', texto_completo, re.IGNORECASE)
    nro_lote = ""
    if m:
        lote_raw = limpiar(m.group(1))
        # Cortar en "Vto", "Vencimiento", "Retest", "Control" si están en la misma línea
        lote_raw = re.split(r'\s+(?:Vto|Vencimiento|Retest|Control)\b', lote_raw, flags=re.IGNORECASE)[0]
        lote_raw = re.sub(r',?\s*enviada\s+el.*$', '', lote_raw, flags=re.IGNORECASE).strip()
        lote_raw = re.sub(r',?\s*a saber.*$', '', lote_raw, flags=re.IGNORECASE).strip()
        # Quitar coma final
        nro_lote = lote_raw.rstrip(',')

    # ── Vencimiento ──
    m = re.search(r'Vencimiento\s*[:.]?\s*(.+?)(?:\n|$)', texto_completo, re.IGNORECASE)
    vencimiento_texto = limpiar(m.group(1)) if m else ""
    # Manejar casos especiales
    if vencimiento_texto.upper() in ('NO FIGURA', 'S/F', 'SIN FECHA', '-'):
        fecha_vencimiento = None
    else:
        fecha_vencimiento = parsear_vencimiento(vencimiento_texto) if vencimiento_texto else None

    # ── IR N° ──
    m = re.search(r'IR\s+N[°º]?\s*[:.]?\s*(\d+/\d+)', texto_completo, re.IGNORECASE)
    nro_ir = m.group(1).strip() if m else ""

    # ── Cantidad remitida o enviada ──
    m = re.search(r'Cantidad\s+(?:remitida|enviada)\s*[:.]?\s*(.+?)(?:\n|$)', texto_completo, re.IGNORECASE)
    cantidad_texto = limpiar(m.group(1)) if m else ""
    stock_inicial = parsear_cantidad(cantidad_texto)

    # ── Unidad de medida y conversión g → mg ──
    unidad = "mg"  # unidad por defecto: miligramos
    if cantidad_texto:
        m_unidad = re.search(r'(\d+(?:[.,]\d+)?)\s*([a-zA-Zμ]+)', cantidad_texto)
        if m_unidad:
            unidad_raw = m_unidad.group(1)
            unidad_doc = m_unidad.group(2).lower().strip()
            if unidad_doc in ('g', 'gr', 'gramos', 'gram', 'grams'):
                # Convertir a miligramos
                unidad = "mg"
                if stock_inicial is not None:
                    stock_inicial = round(stock_inicial * 1000, 4)
            elif unidad_doc in ('mg', 'miligramos', 'milligrams'):
                unidad = "mg"
            elif unidad_doc in ('ml', 'milliliter', 'milliliters', 'mililitros'):
                unidad = "ml"
            else:
                # Cualquier otra unidad (μg, kg, L, etc.) no es válida en LIMSS
                # (solo admite mg/ml) -- se usa mg por defecto en vez de guardar
                # la unidad cruda del documento.
                unidad = "mg"

    # ── Conservación ──
    m = re.search(r'Conservaci[oó]n\s*[:.]?\s*(.+?)(?:\n|$)', texto_completo, re.IGNORECASE)
    conservacion = limpiar(m.group(1)) if m else ""

    # Estrategia 4: formato "St trab" — una línea en negrita con
    # NOMBRE Lote N° XXXX Vto: DD Mes/AA (mayúsculas o minúsculas)
    if not nombre:
        for p in parrafos_a_usar:
            texto_bold = "".join(
                run.text for run in p.runs if run.bold and run.text.strip()
            ).strip()
            texto_bold = limpiar(texto_bold)
            m_trab = re.match(
                r'^(.+?)\s*,?\s*Lote\s+N[°º]',
                texto_bold, re.IGNORECASE
            )
            if m_trab:
                candidato = limpiar(m_trab.group(1))
                # Validar que el candidato sea un nombre real (no un resultado analítico)
                if es_nombre_valido(candidato) and not re.search(r'%|RSD|C\.V|\d{4}-\d{2}-\d{2}', candidato):
                    nombre = candidato
                # También extraer lote y vencimiento de esta línea
                m_lote = re.search(r'Lote\s+N[°º]\s*[:\.]?\s*(.+?)(?:\s+Vto|\s+Vencimiento|$)', texto_bold, re.IGNORECASE)
                if m_lote and not nro_lote:
                    nro_lote = limpiar(m_lote.group(1)).split('Vto')[0].strip()
                m_vto = re.search(r'Vto\s*[:\.]?\s*(.+?)$', texto_bold, re.IGNORECASE)
                if m_vto and not fecha_vencimiento:
                    vto_txt = limpiar(m_vto.group(1))
                    if vto_txt.upper() not in ('NO FIGURA', 'S/F', 'NO TIENE'):
                        fecha_vencimiento = parsear_vencimiento(vto_txt)
                break

    # Estrategia 5: "traceable NOMBRE_PRODUCTO" en negrita
    # Ej: "estándar de referencia secundario traceable PAMOATO DE PIRANTEL"
    if not nombre:
        for p in parrafos_a_usar:
            texto_bold = "".join(
                run.text for run in p.runs if run.bold and run.text.strip()
            ).strip()
            m_trace = re.search(r'traceable\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s\-DE]+?)(?:$|\n)', texto_bold, re.IGNORECASE)
            if m_trace:
                nombre = limpiar(m_trace.group(1))
                break

    # Estrategia 6: primera línea bold que tiene NOMBRE pegado a "Producto Nº:"
    # Ej: "PAMOATO DE PIRANTELProducto Nº: PHR2080-1 G"
    if not nombre:
        for p in parrafos_a_usar:
            texto_bold = "".join(
                run.text for run in p.runs if run.bold and run.text.strip()
            ).strip()
            m_pegado = re.match(r'^([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s\-DE]+?)Producto\s+N[°º]', texto_bold, re.IGNORECASE)
            if m_pegado:
                nombre = limpiar(m_pegado.group(1))
                # Extraer también el producto si no lo tenemos
                m_prod = re.search(r'Producto\s+N[°º]\s*[:\.]?\s*(.+?)(?:\s|$)', texto_bold, re.IGNORECASE)
                if m_prod and not nro_producto:
                    nro_producto = limpiar(m_prod.group(1))
                break

    # Si no hay lote, usar el nro_producto como identificador
    if not nro_lote:
        if nro_producto:
            nro_lote = nro_producto
            logging.warning(f"  Sin Lote Nº — usando Producto Nº como lote: {nro_lote}")
        else:
            logging.error(f"  No se pudo extraer lote ni producto de {docx_path.name}")
            return None

    # Normalizar nombre a mayúsculas
    nombre = nombre.upper().strip()

    return {
        "nombre":            nombre,
        "nro_producto":      nro_producto,
        "nro_lote":          nro_lote,
        "fecha_vencimiento": fecha_vencimiento,
        "nro_ir":            nro_ir,
        "stock_inicial":     stock_inicial if stock_inicial is not None else 0,
        "unidad_medida":     unidad,
        "conservacion":      conservacion,
    }


def es_nombre_valido(texto: str) -> bool:
    """Verifica que el texto sea un nombre de testigo válido y no un lote,
    resultado analítico o línea de log."""
    if not texto or len(texto) < 4:
        return False
    # Descartar si tiene porcentajes o paréntesis de resultados analíticos
    if re.search(r'\d+[,\.]\d+\s*%|RSD|C\.V|Traceable', texto, re.IGNORECASE):
        return False
    # Descartar si parece un lote (mayúsculas + números cortos)
    if re.match(r'^[A-Z]{2,6}\s*\d{4,}$', texto.strip()):
        return False
    # Descartar si parece una línea de log
    if re.search(r'\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}|INFO\s+Producto', texto):
        return False
    # Descartar si empieza con paréntesis o puntuación
    if texto[0] in ('(', ')', ':', '.', '%'):
        return False
    # Debe tener al menos una letra
    if not re.search(r'[A-Za-záéíóúÁÉÍÓÚñÑ]', texto):
        return False
    return True
    cursor.execute(
        "SELECT MAX(CAST(SUBSTRING(codigo, 6, 10) AS INT)) AS max_n FROM lims_testigos WHERE codigo LIKE 'TEST-%'"
    )
    row = cursor.fetchone()
    n = (row.max_n or 0) + 1
    return f"TEST-{n:03d}"


def siguiente_codigo(cursor) -> str:
    cursor.execute(
        "SELECT MAX(CAST(SUBSTRING(codigo, 6, 10) AS INT)) AS max_n FROM lims_testigos WHERE codigo LIKE 'TEST-%'"
    )
    row = cursor.fetchone()
    n = (row.max_n or 0) + 1
    return f"TEST-{n:03d}"


def clave_duplicado(datos: dict) -> str:
    """Clave única: lote + nombre (dos testigos del mismo producto y lote son el mismo)."""
    return f"{datos['nro_lote'].strip().upper()}|{datos['nombre'].strip().upper()}"


def insertar_en_bd(datos: dict, conn) -> bool:
    cursor = conn.cursor()

    # Verificar si ya existe por lote + nombre
    cursor.execute(
        "SELECT id_testigo FROM lims_testigos WHERE nro_lote = ? AND nombre = ?",
        datos["nro_lote"], datos["nombre"]
    )
    if cursor.fetchone():
        logging.info(f"  Ya existe (lote={datos['nro_lote']}) — saltando.")
        return False

    codigo = siguiente_codigo(cursor)

    cursor.execute(
        """
        INSERT INTO lims_testigos
            (codigo, nombre, nro_lote, fecha_vencimiento, stock_actual,
             stock_minimo, unidad_medida, activo, id_usuario_carga)
        VALUES (?, ?, ?, ?, ?, ?, ?, 1, ?)
        """,
        codigo,
        datos["nombre"][:150],
        datos["nro_lote"][:50],
        datos["fecha_vencimiento"].strftime('%Y-%m-%d') if datos["fecha_vencimiento"] else None,
        float(datos["stock_inicial"]),
        float(CONFIG["stock_minimo"]),
        datos["unidad_medida"][:20],
        int(CONFIG["id_usuario_carga"]),
    )
    cursor.execute("SELECT @@IDENTITY AS id")
    id_testigo = int(cursor.fetchone().id)

    conn.commit()
    logging.info(f"  Código asignado: {codigo} (id={id_testigo})")
    return True


def main():
    parser = argparse.ArgumentParser(description="Importar testigos/estándares al LIMSS")
    parser.add_argument("--carpeta", required=True, help="Carpeta raíz (busca recursivamente)")
    parser.add_argument("--dry-run", action="store_true", help="Solo parsear, no insertar en BD")
    args = parser.parse_args()

    carpeta = Path(args.carpeta)
    if not carpeta.exists():
        print(f"ERROR: La carpeta '{carpeta}' no existe.")
        sys.exit(1)

    log_file = setup_logging(carpeta)

    # Buscar .docx recursivamente en subcarpetas
    archivos = sorted([
        f for f in carpeta.rglob("*.docx")
        if not f.name.startswith('~')
    ])
    # También .doc sin convertir
    docs_sin_convertir = [
        f for f in carpeta.rglob("*.doc")
        if not f.name.startswith('~') and not f.with_suffix('.docx').exists()
    ]
    if docs_sin_convertir:
        logging.warning(f"Hay {len(docs_sin_convertir)} archivos .doc sin convertir — se omitirán.")
        logging.warning("Convertirlos primero con LibreOffice o Microsoft Word.")

    logging.info(f"Carpeta raíz: {carpeta}")
    logging.info(f"Archivos .docx encontrados: {len(archivos)}")
    logging.info(f"Modo: {'DRY RUN' if args.dry_run else 'INSERCIÓN EN BD'}")
    logging.info("=" * 60)

    conn = None
    if not args.dry_run:
        try:
            import pyodbc
            cs = (f"DRIVER={{{CONFIG['db_driver']}}};"
                  f"SERVER={CONFIG['db_server']};"
                  f"DATABASE={CONFIG['db_name']};"
                  f"UID={CONFIG['db_user']};"
                  f"PWD={CONFIG['db_password']};"
                  f"TrustServerCertificate=yes;")
            conn = pyodbc.connect(cs, autocommit=False)
            logging.info("✓ Conexión a BD LIMSS OK")
        except Exception as e:
            logging.error(f"✗ No se pudo conectar a la BD: {e}")
            sys.exit(1)

    ok = saltados = errores = 0

    for archivo in archivos:
        logging.info(f"\n→ {archivo.relative_to(carpeta)}")
        datos = parsear_documento(archivo)
        if not datos:
            errores += 1
            continue

        logging.info(f"  Nombre:      {datos['nombre']}")
        logging.info(f"  Producto N°: {datos['nro_producto']}")
        logging.info(f"  Lote:        {datos['nro_lote']}")
        logging.info(f"  Vencimiento: {datos['fecha_vencimiento']}")
        logging.info(f"  IR:          {datos['nro_ir']}")
        logging.info(f"  Stock ini.:  {datos['stock_inicial']} {datos['unidad_medida']}")
        logging.info(f"  Conserv.:    {datos['conservacion']}")

        if args.dry_run:
            ok += 1
            continue

        try:
            if insertar_en_bd(datos, conn):
                ok += 1
            else:
                saltados += 1
        except Exception as e:
            logging.error(f"  ✗ Error al insertar: {e}")
            if conn:
                conn.rollback()
            errores += 1

    if conn:
        conn.close()

    logging.info("\n" + "=" * 60)
    logging.info("RESUMEN FINAL")
    logging.info(f"  ✓ Insertados:  {ok}")
    logging.info(f"  ⏭  Saltados:   {saltados}")
    logging.info(f"  ✗ Errores:     {errores}")
    logging.info("=" * 60)


if __name__ == "__main__":
    main()
