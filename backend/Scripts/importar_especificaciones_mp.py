"""
importar_especificaciones_mp.py
===============================
Importa masivamente los archivos .doc de Especificaciones de Materia Prima
al sistema LIMSS (BD SQL Server).

Requisitos:
    pip install python-docx pyodbc

Uso:
    python importar_especificaciones_mp.py --carpeta "Y:\\Lamar\\Especificaciones\\MP"

Comportamiento:
    - Convierte .doc a .docx con LibreOffice (soffice debe estar instalado)
    - Si LibreOffice no está disponible, intenta leer .docx directamente
    - Solo inserta si el código de MP NO existe en la BD (skip si ya existe)
    - Genera log detallado por archivo y resumen al final
    - El usuario admin debe existir previamente en lims_usuarios

Configuración: editar el bloque CONFIG más abajo.
"""

import os
import sys
import re
import subprocess
import argparse
import logging
from datetime import datetime
from pathlib import Path

# ── CONFIGURACIÓN ──────────────────────────────────────────────
CONFIG = {
    "db_server":   "Lamarserver",
    "db_name":     "LIMSS",
    "db_user":     "limss_app",
    "db_password": "Limss2024#",       # ← cambiar
    "db_driver":   "SQL Server",
    "id_usuario_carga": 1,                   # ← ID del usuario admin en lims_usuarios
    "tipo_material": "materia_prima",
}
# ───────────────────────────────────────────────────────────────


def setup_logging(carpeta_log: Path):
    log_file = carpeta_log / f"importacion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s  %(levelname)-8s  %(message)s",
        handlers=[
            logging.FileHandler(log_file, encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )
    logging.info(f"Log guardado en: {log_file}")
    return log_file


def convertir_doc_a_docx(doc_path: Path) -> Path | None:
    """Convierte .doc a .docx usando LibreOffice o Microsoft Word (si está disponible)."""
    docx_path = doc_path.with_suffix(".docx")
    if docx_path.exists():
        return docx_path

    # Intento 1: LibreOffice
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx",
             "--outdir", str(doc_path.parent), str(doc_path)],
            capture_output=True, timeout=30
        )
        if result.returncode == 0 and docx_path.exists():
            return docx_path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Intento 2: Microsoft Word via COM (solo Windows con Word instalado)
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(doc_path.resolve()))
        doc.SaveAs2(str(docx_path.resolve()), FileFormat=16)  # 16 = docx
        doc.Close()
        word.Quit()
        if docx_path.exists():
            logging.info(f"  Convertido con Microsoft Word OK")
            return docx_path
    except Exception as e:
        logging.warning(f"  Microsoft Word no disponible: {e}")

    logging.error(f"  No se pudo convertir {doc_path.name} — instalar LibreOffice o Microsoft Word")
    return None


def limpiar(texto: str) -> str:
    """Limpia espacios y caracteres extraños."""
    if not texto:
        return ""
    return " ".join(texto.strip().split())


def extraer_campo_lista(texto: str, campo: str) -> str:
    """Extrae el valor de una línea con patrón '- Campo: valor'."""
    patron = rf"[-•]\s*{re.escape(campo)}\s*:\s*(.+)"
    match = re.search(patron, texto, re.IGNORECASE)
    if match:
        return limpiar(match.group(1).strip())
    return ""


def parsear_documento(docx_path: Path) -> dict | None:
    """Parsea el .docx y extrae los datos de secciones 1 y 2."""
    try:
        from docx import Document
    except ImportError:
        logging.error("  python-docx no instalado. Ejecutar: pip install python-docx")
        return None

    try:
        doc = Document(docx_path)
    except Exception as e:
        logging.error(f"  No se pudo abrir {docx_path.name}: {e}")
        return None

    # Extraer texto completo por párrafos y tablas del cuerpo
    texto_completo = []
    tablas_ensayos = []  # guardar tablas para procesar ensayos después

    for element in doc.element.body:
        from docx.oxml.ns import qn
        if element.tag == qn('w:p'):
            from docx.text.paragraph import Paragraph
            p = Paragraph(element, doc)
            if p.text.strip():
                texto_completo.append(p.text.strip())
        elif element.tag == qn('w:tbl'):
            from docx.table import Table
            t = Table(element, doc)
            rows_data = []
            for row in t.rows:
                row_texts = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_texts)
                texto_completo.append(" | ".join(row_texts))
            tablas_ensayos.append(rows_data)

    texto = "\n".join(texto_completo)

    # ── Leer el header del documento (donde está el código MP) ──
    header_texts = []
    for section in doc.sections:
        header = section.header
        for t in header.tables:
            for row in t.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        header_texts.append(txt)
        for p in header.paragraphs:
            if p.text.strip():
                header_texts.append(p.text.strip())

    # ── Código de MP (del header) ──
    codigo = ""
    nombre = ""
    tipo_mp = ""
    accion = ""

    for txt in header_texts:
        txt_clean = txt.replace('\n', ' ').strip()
        # Código: MP078-00 → tomar solo MP078
        m = re.search(r'C[oó]digo\s*:\s*(MP\d{3})', txt_clean, re.IGNORECASE)
        if m:
            codigo = m.group(1).strip()
        # Nombre del producto
        m2 = re.search(r'Especificaciones de Materia Prima\s*(.+)', txt_clean, re.IGNORECASE | re.DOTALL)
        if m2:
            nombre = limpiar(m2.group(1).replace('\n', ' '))
        # Tipo
        m3 = re.search(r'Tipo\s*:\s*(.+)', txt_clean, re.IGNORECASE)
        if m3:
            tipo_mp = limpiar(m3.group(1))
        # Acción terapéutica
        m4 = re.search(r'Acci[oó]n Terap[eé]utica\s*:\s*(.+)', txt_clean, re.IGNORECASE)
        if m4:
            accion = limpiar(m4.group(1))

    if not codigo:
        logging.error(f"  No se pudo extraer el código de MP de {docx_path.name}")
        return None

    # ── Sección 1: Información ──
    sinonimia = extraer_campo_lista(texto, "Sinonimia")
    nro_cas   = extraer_campo_lista(texto, r"CAS\s*N[°º]")
    # Limpiar corchetes del CAS
    nro_cas   = nro_cas.strip("[]").strip()

    # Fórmula molecular
    fm_match = re.search(r'F[oó]rmula\s+molecular.*?(?:FM\s*)?:\s*(.+?)(?:\n|$)', texto, re.IGNORECASE)
    formula_molecular = limpiar(fm_match.group(1)) if fm_match else ""

    # Peso molecular (puede tener "Anhidro: X  Monohidrato: Y" — guardamos todo)
    pm_match = re.search(r'Peso\s+molecular.*?(?:PM\s*)?:\s*(.+?)(?:\n- |$)', texto, re.IGNORECASE | re.DOTALL)
    peso_molecular_texto = limpiar(pm_match.group(1).replace('\n', ' ')) if pm_match else ""

    # Nombre químico (puede tener variantes)
    nq_match = re.search(r'Nombre\s+qu[ií]mico\s*:\s*(.*?)(?=\n-\s*F[oó]rmula|\n-\s*Peso)', texto, re.IGNORECASE | re.DOTALL)
    nombre_quimico = limpiar(nq_match.group(1).replace('\n', ' ')) if nq_match else ""

    # Envasado
    env_match = re.search(r'Envasado\s+y\s+almacenamiento\s*:\s*(.+?)(?:\n#|\n- |\Z)', texto, re.IGNORECASE | re.DOTALL)
    envasado = limpiar(env_match.group(1).replace('\n', ' ')) if env_match else ""

    # ── Sección 2: Ensayos (de las tablas reales del documento) ──
    ensayos = []
    HEADERS_ENSAYO = {'ensayo', 'determinación', 'determinacion'}

    for tabla in tablas_ensayos:
        if not tabla:
            continue
        # Verificar si la primera fila es encabezado de ensayos
        primera_fila = [c.lower().strip() for c in tabla[0]]
        if not any(h in ' '.join(primera_fila) for h in HEADERS_ENSAYO):
            continue
        # Procesar filas de datos
        for fila in tabla[1:]:
            if len(fila) < 2:
                continue
            nombre_ensayo = limpiar(fila[0])
            especificacion = limpiar(fila[1]) if len(fila) > 1 else ""
            tecnica = limpiar(fila[2]) if len(fila) > 2 else ""
            bibliografia = limpiar(fila[3]) if len(fila) > 3 else ""
            if nombre_ensayo and not re.match(r'^(ENSAYO|DETERMINACI[OÓ]N)$', nombre_ensayo, re.IGNORECASE):
                ensayos.append({
                    "nombre_ensayo":  nombre_ensayo,
                    "especificacion": especificacion,
                    "tecnica":        tecnica,
                    "bibliografia":   bibliografia,
                })

    return {
        "codigo":             codigo,
        "nombre":             nombre,
        "tipo_mp":            tipo_mp,
        "accion_terapeutica": accion,
        "sinonimia":          sinonimia,
        "nro_cas":            nro_cas,
        "nombre_quimico":     nombre_quimico,
        "formula_molecular":  formula_molecular,
        "peso_molecular":     peso_molecular_texto,
        "envasado_almacenamiento": envasado,
        "ensayos":            ensayos,
    }


def insertar_en_bd(datos: dict, conn) -> bool:
    """Inserta la especificación y sus ensayos en LIMSS. Retorna True si insertó."""
    cursor = conn.cursor()

    # Verificar si ya existe
    cursor.execute(
        "SELECT id_especificacion FROM lims_especificaciones WHERE erp_CODART = ?",
        datos["codigo"]
    )
    if cursor.fetchone():
        logging.info(f"  ⏭  Ya existe en BD — saltando.")
        return False

    # Debug: verificar largos antes de insertar
    campos_spec = {
        "erp_CODART": (datos["codigo"], 20),
        "erp_DESART": (datos["nombre"], 100),
        "tipo_mp": (datos["tipo_mp"], 50),
        "accion_terapeutica": (datos["accion_terapeutica"], 500),
        "sinonimia": (datos["sinonimia"], 1000),
        "nro_cas": (datos["nro_cas"], 50),
        "nombre_quimico": (datos["nombre_quimico"], 2000),
        "formula_molecular": (datos["formula_molecular"], 100),
        "peso_molecular": (datos["peso_molecular"], 100),
        "envasado_almacenamiento": (datos["envasado_almacenamiento"], 2000),
    }
    for campo, (valor, limite) in campos_spec.items():
        if valor and len(valor) > limite:
            logging.warning(f"  TRUNCAMIENTO en {campo}: len={len(valor)} > {limite} — '{valor[:80]}...'")

    # Insertar especificación — truncar preventivamente campos VARCHAR
    cursor.execute(
        """
        INSERT INTO lims_especificaciones
            (erp_IdM21, erp_CODART, erp_DESART, tipo_material, version, vigente,
             tipo_mp, accion_terapeutica, sinonimia, nro_cas, nombre_quimico,
             formula_molecular, peso_molecular, envasado_almacenamiento,
             id_usuario_carga)
        VALUES (0, ?, ?, ?, '1.0', 1, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        datos["codigo"][:20],
        datos["nombre"][:100],
        CONFIG["tipo_material"][:20],
        datos["tipo_mp"][:50],
        datos["accion_terapeutica"][:500],
        datos["sinonimia"][:1000],
        datos["nro_cas"][:50],
        datos["nombre_quimico"][:2000],
        datos["formula_molecular"][:100],
        datos["peso_molecular"][:100],
        datos["envasado_almacenamiento"][:2000],
        CONFIG["id_usuario_carga"],
    )
    cursor.execute("SELECT @@IDENTITY AS id")
    id_especificacion = int(cursor.fetchone().id)

    # Insertar ensayos
    for orden, ensayo in enumerate(datos["ensayos"], start=1):
        campos_ensayo = {
            "nombre_ensayo": (ensayo["nombre_ensayo"], 300),
            "metodologia": (ensayo["tecnica"], 500),
            "valor_requerido": (ensayo["especificacion"], 2000),
            "bibliografia": (ensayo["bibliografia"], 500),
            "especificacion_texto": (ensayo["especificacion"], 2000),
        }
        for campo, (valor, limite) in campos_ensayo.items():
            if valor and len(valor) > limite:
                logging.warning(f"  TRUNCAMIENTO ensayo '{ensayo['nombre_ensayo']}' campo {campo}: len={len(valor)} > {limite}")
        cursor.execute(
            """
            INSERT INTO lims_ensayos
                (id_especificacion, orden, nombre_ensayo, metodologia,
                 tipo_dato, valor_requerido, bibliografia, especificacion_texto, obligatorio)
            VALUES (?, ?, ?, ?, 'cualitativo', ?, ?, ?, 1)
            """,
            id_especificacion,
            orden,
            ensayo["nombre_ensayo"][:300],
            ensayo["tecnica"][:500],
            ensayo["especificacion"][:2000],
            ensayo["bibliografia"][:500],
            ensayo["especificacion"][:2000],
        )

    conn.commit()
    return True


def main():
    parser = argparse.ArgumentParser(description="Importar especificaciones MP al LIMSS")
    parser.add_argument("--carpeta", required=True, help="Carpeta con los archivos .doc")
    parser.add_argument("--dry-run", action="store_true", help="Solo parsear, no insertar en BD")
    args = parser.parse_args()

    carpeta = Path(args.carpeta)
    if not carpeta.exists():
        print(f"ERROR: La carpeta '{carpeta}' no existe.")
        sys.exit(1)

    log_file = setup_logging(carpeta)

    # Buscar todos los .doc y .docx
    archivos = sorted(list(carpeta.glob("*.doc")) + list(carpeta.glob("*.docx")))
    archivos = [f for f in archivos if not f.name.endswith(".docx") or
                not (f.with_suffix(".doc")).exists()]  # evitar duplicados

    logging.info(f"Carpeta: {carpeta}")
    logging.info(f"Archivos encontrados: {len(archivos)}")
    logging.info(f"Modo: {'DRY RUN (sin insertar)' if args.dry_run else 'INSERCIÓN EN BD'}")
    logging.info("=" * 60)

    # Conectar a BD
    conn = None
    if not args.dry_run:
        try:
            import pyodbc
            cs = (f"DRIVER={{{CONFIG['db_driver']}}};"
                  f"SERVER={CONFIG['db_server']};"
                  f"DATABASE={CONFIG['db_name']};"
                  f"UID={CONFIG['db_user']};"
                  f"PWD={CONFIG['db_password']};"
                  f"TrustServerCertificate=yes;")
            conn = pyodbc.connect(cs, autocommit=False)
            conn.cursor().execute("SET NOCOUNT OFF")
            logging.info("✓ Conexión a BD LIMSS OK")
        except Exception as e:
            logging.error(f"✗ No se pudo conectar a la BD: {e}")
            sys.exit(1)

    # Contadores
    ok = 0
    saltados = 0
    errores = 0

    for archivo in archivos:
        logging.info(f"\n→ {archivo.name}")

        # Convertir si es .doc
        if archivo.suffix.lower() == ".doc":
            docx = convertir_doc_a_docx(archivo)
            if not docx:
                logging.error(f"  ✗ No se pudo convertir a .docx")
                errores += 1
                continue
        else:
            docx = archivo

        # Parsear
        datos = parsear_documento(docx)
        if not datos:
            logging.error(f"  ✗ Error al parsear el documento")
            errores += 1
            continue

        logging.info(f"  Código:  {datos['codigo']}")
        logging.info(f"  Nombre:  {datos['nombre']}")
        logging.info(f"  Tipo:    {datos['tipo_mp']}")
        logging.info(f"  CAS:     {datos['nro_cas']}")
        logging.info(f"  Ensayos: {len(datos['ensayos'])}")
        for e in datos["ensayos"]:
            logging.info(f"    • {e['nombre_ensayo']}")

        if args.dry_run:
            logging.info(f"  [DRY RUN] No se insertó en BD")
            ok += 1
            continue

        # Insertar
        try:
            insertado = insertar_en_bd(datos, conn)
            if insertado:
                logging.info(f"  ✓ Insertado correctamente")
                ok += 1
            else:
                saltados += 1
        except Exception as e:
            logging.error(f"  ✗ Error al insertar: {e}")
            # Mostrar largos de todos los campos para diagnóstico
            logging.error(f"    Diagnóstico de largos:")
            logging.error(f"    erp_CODART={len(datos['codigo'])} erp_DESART={len(datos['nombre'])} tipo_mp={len(datos['tipo_mp'])} accion={len(datos['accion_terapeutica'])}")
            for i, ens in enumerate(datos['ensayos']):
                logging.error(f"    Ensayo {i+1}: nombre={len(ens['nombre_ensayo'])} espec={len(ens['especificacion'])} tecnica={len(ens['tecnica'])} biblio={len(ens['bibliografia'])}")
            conn.rollback()
            errores += 1

    if conn:
        conn.close()

    # Resumen final
    logging.info("\n" + "=" * 60)
    logging.info("RESUMEN FINAL")
    logging.info(f"  Total archivos procesados: {len(archivos)}")
    logging.info(f"  ✓ Insertados correctamente: {ok}")
    logging.info(f"  ⏭  Saltados (ya existían):  {saltados}")
    logging.info(f"  ✗ Errores:                 {errores}")
    logging.info(f"  Log guardado en:            {log_file}")
    logging.info("=" * 60)


if __name__ == "__main__":
    main()
